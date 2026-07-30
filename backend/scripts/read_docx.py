import docx
import sys

def print_docx(filename):
    try:
        doc = docx.Document(filename)
        for p in doc.paragraphs:
            style_name = p.style.name if p.style else "No Style"
            if p.text.strip(): print(f"Style: {style_name} | Text: {p.text}")
        for t in doc.tables:
            print("--- TABLE ---")
            for row in t.rows:
                row_data = [cell.text.replace('\n', ' ') for cell in row.cells]
                print(" | ".join(row_data).encode('ascii', 'ignore').decode('ascii'))
    except Exception as e:
        print(e)

if __name__ == '__main__':
    print_docx(sys.argv[1])
